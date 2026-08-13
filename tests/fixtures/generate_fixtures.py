"""Generates tests/fixtures/sample_syllabus.docx — a synthetic but realistic
course syllabus used by several tests instead of depending on a file outside
the repository. Re-run this script if the fixture needs regenerating:

    python tests/fixtures/generate_fixtures.py

Content requirements baked in on purpose (don't remove these phrases without
checking which tests rely on them):
  - "Absolutely no late work" / "Late Work" section  -> retrieval + late-work-policy tests
  - "Required Text"                                   -> synonym-expansion test
  - total length comfortably over 21,000 characters    -> chunking tests (>5 chunks)
"""
import os

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

OUT_DOCX = os.path.join(os.path.dirname(__file__), "sample_syllabus.docx")
OUT_PDF = os.path.join(os.path.dirname(__file__), "sample_cs_syllabus.pdf")


def build_docx():
    doc = Document()
    doc.add_heading("CIS 3305 — Systems Analysis and Design", level=1)
    doc.add_paragraph("Spring 2026 | CRN 12345 | 3 Credit Hours")
    doc.add_paragraph("Instructor: Dr. A. Rivera  |  Email: arivera@example.edu  |  Office: Bell Hall 302")

    doc.add_heading("Course Description", level=2)
    doc.add_paragraph(
        "This course introduces the principles and practices of systems analysis and design "
        "used to build modern information systems. Topics include requirements elicitation, "
        "process modeling, data modeling, use-case analysis, system architecture, and an "
        "introduction to agile development practices. Students will work individually and in "
        "small teams to analyze a real-world business problem and produce a complete system "
        "design proposal by the end of the semester."
    )

    doc.add_heading("Required Text", level=2)
    doc.add_paragraph(
        "Required Text: Systems Analysis and Design, 12th Edition, by Kendall & Kendall. "
        "ISBN 978-0-13-489333-4. A digital copy is available through the campus bookstore "
        "and is also on 2-hour reserve at the university library. Additional readings will be "
        "posted to the course page throughout the semester and are also required."
    )

    doc.add_heading("Learning Outcomes", level=2)
    for i, outcome in enumerate([
        "Elicit and document functional and non-functional requirements from stakeholders.",
        "Construct use-case diagrams, activity diagrams, and entity-relationship diagrams.",
        "Evaluate trade-offs between competing system architectures.",
        "Apply agile principles to iterative system design.",
        "Communicate technical design decisions to non-technical stakeholders.",
    ], 1):
        doc.add_paragraph(f"{i}. {outcome}")

    doc.add_heading("Grading Breakdown", level=2)
    doc.add_paragraph("Homework: 20%")
    doc.add_paragraph("Midterm Exam: 25%")
    doc.add_paragraph("Final Exam: 25%")
    doc.add_paragraph("Team Design Project: 20%")
    doc.add_paragraph("Participation: 10%")

    doc.add_heading("Late Work Policy", level=2)
    doc.add_paragraph(
        "Absolutely no late work will be accepted for homework assignments except in the case "
        "of a documented medical or family emergency, approved in advance whenever possible. "
        "Assignments submitted after the posted deadline will receive a grade of zero. If you "
        "anticipate a conflict with a deadline, contact the instructor at least 48 hours before "
        "the assignment is due to discuss options. Extensions are granted at the instructor's "
        "discretion and are not guaranteed. The team design project has a separate late policy "
        "described in the project handout, since late team submissions affect every member of "
        "the team and are handled with additional care."
    )

    doc.add_heading("Academic Integrity", level=2)
    doc.add_paragraph(
        "Students are expected to comply with the university's academic integrity policy in "
        "all coursework. Collaboration on homework is permitted unless explicitly stated "
        "otherwise, but all submitted work must represent the student's own understanding. "
        "Suspected violations will be referred to the Office of Student Conduct."
    )

    doc.add_heading("Weekly Schedule", level=2)
    topics = [
        "Course Introduction & the Systems Development Life Cycle",
        "Requirements Elicitation Techniques",
        "Business Process Modeling",
        "Use-Case Analysis and Diagrams",
        "Data Flow Diagrams",
        "Entity-Relationship Modeling",
        "Normalization and Database Design",
        "Midterm Exam Review",
        "Midterm Exam",
        "System Architecture Patterns",
        "User Interface Design Principles",
        "Agile and Iterative Development",
        "Team Project Work Sessions",
        "Testing and Quality Assurance",
        "Deployment and Change Management",
        "Team Project Presentations",
        "Final Exam Review",
        "Final Exam",
    ]
    for i, topic in enumerate(topics, 1):
        doc.add_paragraph(
            f"Week {i}: {topic}. Readings assigned for this week cover the corresponding "
            f"chapter in the required text along with any supplementary material posted to "
            f"the course page. Students should come to class prepared to discuss the assigned "
            f"readings and participate in in-class exercises related to {topic.lower()}."
        )

    doc.add_heading("Office Hours and Communication", level=2)
    doc.add_paragraph(
        "Office hours are held Tuesdays and Thursdays from 1:00 PM to 3:00 PM in Bell Hall 302, "
        "or by appointment. Email is the best way to reach the instructor outside of office "
        "hours; please allow up to 48 hours for a response, longer over weekends and holidays. "
        "Course announcements will be posted to the learning management system and students "
        "are responsible for checking it regularly throughout the semester."
    )

    doc.add_heading("Accommodations", level=2)
    doc.add_paragraph(
        "Students requiring accommodations for a documented disability should contact the "
        "campus disability services office as early in the semester as possible so that "
        "arrangements can be made in a timely manner. Accommodations cannot be applied "
        "retroactively to work already submitted or exams already taken."
    )

    doc.add_heading("Team Design Project — Detailed Description", level=2)
    doc.add_paragraph(
        "Working in teams of three to four, students will select a real or realistic business "
        "problem and produce a complete systems analysis and design proposal over the course "
        "of the semester. The project unfolds in four phases, each building on the deliverable "
        "before it, and each phase is graded both on the quality of the artifact produced and "
        "on the team's documented process for producing it."
    )
    doc.add_paragraph(
        "Phase 1 — Requirements Elicitation: teams identify stakeholders, conduct at least two "
        "mock interviews, and produce a requirements document distinguishing functional from "
        "non-functional requirements. This phase emphasizes asking good questions over jumping "
        "to a solution; teams that propose a design before requirements are fully understood "
        "will be asked to revise before proceeding to Phase 2."
    )
    doc.add_paragraph(
        "Phase 2 — Process and Data Modeling: teams produce use-case diagrams, at least one "
        "detailed use-case description, a data flow diagram, and an entity-relationship diagram "
        "for the proposed system. Diagrams should follow standard UML or DFD notation as covered "
        "in lecture, and every entity in the ER diagram must trace back to a requirement from "
        "Phase 1."
    )
    doc.add_paragraph(
        "Phase 3 — Architecture and Interface Design: teams propose a system architecture "
        "(e.g. client-server, layered, microservice) with justification for the trade-offs "
        "involved, and produce wireframes or mockups for the system's primary user interfaces. "
        "Teams should be able to explain why their chosen architecture fits the non-functional "
        "requirements identified earlier, not just the functional ones."
    )
    doc.add_paragraph(
        "Phase 4 — Final Presentation and Report: teams deliver a 15-minute presentation to the "
        "class summarizing the full design, followed by questions from the instructor and peers, "
        "and submit a final written report consolidating all four phases into a single coherent "
        "document. The final report should read as a standalone artifact a real stakeholder could "
        "hand off to a development team."
    )

    doc.add_heading("Rubric Overview for the Team Project", level=2)
    doc.add_paragraph(
        "Each phase is scored across four criteria: completeness (does the deliverable cover "
        "everything the phase requires), correctness (are the diagrams and documents technically "
        "sound), traceability (can every design decision be traced back to a requirement), and "
        "clarity (could someone outside the team understand the deliverable without additional "
        "explanation). The full rubric with point breakdowns for each criterion is posted "
        "separately on the course page and reviewed in detail during the Phase 1 work session."
    )

    doc.add_heading("Classroom Expectations", level=2)
    doc.add_paragraph(
        "Class meets twice weekly and combines short lectures with in-class exercises, so regular "
        "attendance matters more in this course than in a purely lecture-based class. Laptops are "
        "permitted for note-taking and in-class modeling exercises but should otherwise remain "
        "closed during lecture segments. Respectful disagreement during design critiques is "
        "encouraged — a large part of systems analysis is learning to defend a design decision "
        "and to update it gracefully when a critique reveals a real gap."
    )

    doc.add_heading("Exam Format", level=2)
    doc.add_paragraph(
        "Both the midterm and final exams combine short-answer conceptual questions with a "
        "practical modeling component, where students are given a short case description and "
        "asked to produce part of a diagram (e.g. a use-case diagram or an ER diagram) under "
        "exam conditions. The final exam is not comprehensive in the sense of re-testing every "
        "midterm topic in depth, but foundational concepts from the first half of the course are "
        "assumed knowledge for the second half and may appear as supporting context in final "
        "exam questions."
    )

    doc.add_heading("Tools Used in This Course", level=2)
    doc.add_paragraph(
        "Diagramming exercises will use a free, browser-based diagramming tool introduced in "
        "Week 2 — no paid software is required for this course. Team project submissions should "
        "be in PDF format unless a phase explicitly calls for an editable diagram file. Students "
        "are welcome to use whatever note-taking or project-management tools work best for their "
        "team, but the instructor will only be able to provide direct support for the tools "
        "introduced during lecture."
    )

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)

    total_chars = sum(len(p.text) for p in doc.paragraphs)
    print(f"Wrote {OUT_DOCX} (~{total_chars} chars across paragraphs)")


def build_pdf():
    """A genuine 2-page PDF (ABET-style CS syllabus) with real extractable
    text — used by the PDF-page-limit tests, which specifically need a
    document with MORE than one page."""
    c = canvas.Canvas(OUT_PDF, pagesize=LETTER)
    width, height = LETTER
    margin = 1 * inch
    line_height = 14

    page1_lines = [
        "CS 2302 — Data Structures  |  ABET Course Syllabus  |  Spring 2026",
        "CRN 27062  |  3 Credit Hours  |  Prerequisite: CS 1302",
        "",
        "Catalog Description:",
        "Abstract data types, algorithm analysis, and their implementation using",
        "arrays, linked lists, stacks, queues, trees, and graphs. Emphasis on",
        "choosing appropriate data structures and analyzing time/space complexity.",
        "",
        "Course Learning Outcomes (mapped to ABET Student Outcomes):",
        "1. Analyze the asymptotic time and space complexity of an algorithm.",
        "2. Implement and evaluate linear data structures (arrays, linked lists,",
        "   stacks, and queues) for a given problem.",
        "3. Implement and evaluate tree-based and graph-based data structures.",
        "4. Select an appropriate data structure given functional and performance",
        "   requirements of a problem.",
        "5. Communicate algorithmic trade-offs clearly in written form.",
        "",
        "Textbook:",
        "Data Structures and Algorithm Analysis, 3rd Edition, Mark Allen Weiss.",
        "",
        "Grading:",
        "Homework 25%, Programming Projects 30%, Midterm 20%, Final Exam 25%.",
    ]

    page2_lines = [
        "CS 2302 — Course Policies (page 2 of 2)",
        "",
        "Attendance: Attendance is not directly graded but is strongly correlated",
        "with performance on programming projects, which build on each other",
        "week to week.",
        "",
        "Programming Projects: All programming projects must compile and run in",
        "the department's standard Linux lab environment. Projects that do not",
        "compile receive a maximum of 50% credit, regardless of partial",
        "correctness, so test in the lab environment before submitting.",
        "",
        "Academic Integrity: Programming projects are individual work unless a",
        "project description explicitly states otherwise. Submitting code you",
        "did not write, including AI-generated code presented as your own work,",
        "is a violation of the university's academic integrity policy.",
        "",
        "ABET Assessment Note: Select assignments in this course are used for",
        "ABET program assessment purposes. Your work may be reviewed in",
        "de-identified form as part of the department's continuous improvement",
        "process; this does not affect your grade.",
        "",
        "Office Hours: Mondays and Wednesdays, 2:00-3:30 PM, Chemistry & "
        "Computer Science Building, Room 214, or by appointment.",
    ]

    def draw_page(lines):
        c.setFont("Helvetica", 11)
        y = height - margin
        for line in lines:
            c.drawString(margin, y, line)
            y -= line_height
        c.showPage()

    draw_page(page1_lines)
    draw_page(page2_lines)
    c.save()
    print(f"Wrote {OUT_PDF} (2 pages)")


if __name__ == "__main__":
    build_docx()
    build_pdf()
